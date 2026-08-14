from pathlib import Path
from typing import Dict, List
from docx import Document
from .base import BaseParser, ParsedDocument

class DocxParser(BaseParser):
    def supported_extensions(self) -> List[str]:
        return [".docx"]
    
    def parse(self, file_path: Path) -> ParsedDocument:
        doc = Document(str(file_path))
        result = ParsedDocument()
        
        for para in doc.paragraphs[:5]:
            if para.style.name.startswith('Heading') or (para.runs and para.runs[0].bold):
                result.title = para.text.strip()
                break
        
        if not result.title:
            result.title = doc.paragraphs[0].text.strip() if doc.paragraphs else "未命名"
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text.strip(),
                    "style": para.style.name,
                    "level": self._get_heading_level(para.style.name)
                })
        
        result.content = "\n".join([p["text"] for p in paragraphs])
        result.sections = paragraphs
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            result.tables.append(table_data)
        
        return result
    
    def _get_heading_level(self, style_name: str) -> int:
        if style_name.startswith('Heading'):
            try:
                return int(style_name.replace('Heading', ''))
            except:
                return 0
        return 0
