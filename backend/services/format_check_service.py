# -*- coding: utf-8 -*-
"""公文格式校验服务（完整文件，直接覆盖 backend/services/format_check_service.py）

设计要点：
1. 程序规则校验（确定性）：用 python-docx 读取 Word 的字体/字号/对齐/行距/缩进/页边距等
   真实排版属性，与 FormatRule 中配置的期望值逐项比对。
2. AI 辅助判断：只处理规则难以表达的问题（落款规范、结构完整性、明显异常），
   结果与规则校验统一格式返回，source 字段区分 "rule" / "ai"。
3. 自动修正：apply_fixes 按 issue 中的 fix_hint 和 paragraph_index 回写 docx，
   供"审阅模式"的修正预览与修正稿下载使用。
4. 规则不写死：所有期望值来自数据库 FormatRule，司法局正式规范确定后仅需在后台录入。
"""
import re
import json as _json
from pathlib import Path
from typing import List, Dict, Optional

from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.config.settings import FORMAT_CHECK_AI_BUDGET

# 对齐方式映射
_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_ALIGN_NAME = {v: k for k, v in _ALIGN_MAP.items()}
_ALIGN_CN = {"left": "左对齐", "center": "居中", "right": "右对齐", "justify": "两端对齐"}

# 公文标题层级识别
_RE_H1 = re.compile(r'^[一二三四五六七八九十]+、')
_RE_H2 = re.compile(r'^（[一二三四五六七八九十]+）')
_RE_H3 = re.compile(r'^\d+[\.、]')


def _run_font_name(run) -> Optional[str]:
    """读取 run 的有效中文字体名（优先 eastAsia）"""
    rpr = run._element.rPr
    if rpr is not None:
        rfonts = rpr.rFonts
        if rfonts is not None:
            ea = rfonts.get(qn('w:eastAsia'))
            if ea:
                return ea
            ascii_f = rfonts.get(qn('w:ascii'))
            if ascii_f:
                return ascii_f
    return run.font.name


def _fmt_value(kind: str, value) -> str:
    """把属性值格式化为人类可读描述"""
    if value is None:
        return "未设置"
    if kind == "font_size_pt":
        return f"{value:g}磅"
    if kind == "alignment":
        return _ALIGN_CN.get(value, str(value))
    if kind == "bold":
        return "加粗" if value else "不加粗"
    if kind.endswith("_cm"):
        return f"{value:g}厘米"
    if kind.endswith("_pt"):
        return f"{value:g}磅"
    if kind == "first_line_indent_chars":
        return f"缩进{value}字符"
    return str(value)


class FormatCheckService:
    """格式校验引擎：规则校验 + AI 辅助 + 自动修正"""

    def __init__(self, llm_service=None):
        self.llm = llm_service

    # ================= 主入口 =================
    def check_file(self, file_path: Path, filename: str, rules: List[Dict],
                   use_ai: bool = True) -> Dict:
        suffix = filename.lower().rsplit('.', 1)[-1]
        if suffix == "docx":
            return self._check_docx(file_path, filename, rules, use_ai)
        elif suffix in ("txt", "md"):
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            return self._check_plain_text(text, filename, rules, use_ai)
        elif suffix == "pdf":
            return {
                "filename": filename,
                "file_type": suffix,
                "issues": [{
                    "location": "全文",
                    "element": "文件类型",
                    "current": "PDF 文件",
                    "expected": "Word (.docx) 文件",
                    "suggestion": "PDF 不包含字体/字号等排版属性，无法进行精确格式校验。"
                                "建议转换为 Word 后再校验；如需检查 PDF 内容规范性，可使用 AI 辅助分析。",
                    "source": "rule", "severity": "warning",
                    "paragraph_index": None, "fix_hint": None
                }],
                "ai_used": False
            }
        else:
            raise ValueError(f"格式校验暂不支持 .{suffix} 文件，请上传 .docx 文件")

    # ================= docx 规则校验 =================
    def _check_docx(self, file_path: Path, filename: str, rules: List[Dict],
                    use_ai: bool) -> Dict:
        doc = DocxDocument(str(file_path))
        issues: List[Dict] = []

        rule_map: Dict[str, List[Dict]] = {}
        for r in rules:
            rule_map.setdefault(r["target"], []).append(r)

        paragraphs = [p for p in doc.paragraphs]
        nonempty = [p for p in paragraphs if p.text.strip()]

        # ---- 页面设置 ----
        for rule in rule_map.get("page", []):
            issues.extend(self._check_page(doc, rule))

        # ---- 标题（第一个非空段落视为标题） ----
        title_para = nonempty[0] if nonempty else None
        for rule in rule_map.get("title", []):
            if title_para is not None:
                issues.extend(self._check_paragraph(title_para, 0, rule, "第1段（标题）"))

        # ---- 正文 / 标题层级 ----
        body_rules = rule_map.get("body", [])
        h1_rules = rule_map.get("heading1", [])
        h2_rules = rule_map.get("heading2", [])
        for idx, p in enumerate(paragraphs):
            text = p.text.strip()
            if not text:
                continue
            if title_para is not None and p is title_para:
                continue
            label = f"第{idx + 1}段"
            if _RE_H1.match(text):
                for rule in h1_rules:
                    issues.extend(self._check_paragraph(p, idx, rule, f"{label}（一级标题）"))
            elif _RE_H2.match(text):
                for rule in h2_rules:
                    issues.extend(self._check_paragraph(p, idx, rule, f"{label}（二级标题）"))
            else:
                for rule in body_rules:
                    issues.extend(self._check_paragraph(p, idx, rule, label))

        # ---- 落款 / 日期（文末段落） ----
        sig_rules = rule_map.get("signature", [])
        date_rules = rule_map.get("date", [])
        if sig_rules or date_rules:
            tail = nonempty[-2:] if len(nonempty) >= 2 else nonempty[-1:]
            for p in tail:
                idx = paragraphs.index(p)
                text = p.text.strip()
                is_date = bool(re.search(r'\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日', text))
                target_rules = date_rules if is_date else sig_rules
                kind = "成文日期" if is_date else "落款"
                for rule in target_rules:
                    issues.extend(self._check_paragraph(p, idx, rule, f"文末（{kind}）"))

        # ---- 全文通用检查（空行、行尾空格等） ----
        for rule in rule_map.get("general", []):
            issues.extend(self._check_general(paragraphs, rule))

        # ---- AI 辅助 ----
        ai_used = False
        if use_ai and self.llm is not None:
            full_text = "\n".join(p.text for p in paragraphs if p.text.strip())
            ai_issues = self._ai_check(full_text[:FORMAT_CHECK_AI_BUDGET], rules)
            if ai_issues is not None:
                issues.extend(ai_issues)
                ai_used = True

        return {
            "filename": filename,
            "file_type": "docx",
            "issues": issues,
            "ai_used": ai_used
        }

    # ---- 页面 ----
    def _check_page(self, doc, rule: Dict) -> List[Dict]:
        issues = []
        checks = rule.get("checks", {})
        section = doc.sections[0]
        current = {
            "top_margin_cm": section.top_margin.cm if section.top_margin else None,
            "bottom_margin_cm": section.bottom_margin.cm if section.bottom_margin else None,
            "left_margin_cm": section.left_margin.cm if section.left_margin else None,
            "right_margin_cm": section.right_margin.cm if section.right_margin else None,
            "page_width_cm": section.page_width.cm if section.page_width else None,
            "page_height_cm": section.page_height.cm if section.page_height else None,
        }
        names = {
            "top_margin_cm": "上边距", "bottom_margin_cm": "下边距",
            "left_margin_cm": "左边距", "right_margin_cm": "右边距",
            "page_width_cm": "页面宽度", "page_height_cm": "页面高度",
        }
        for key, expected in checks.items():
            if key not in current or expected is None:
                continue
            cur = current[key]
            if cur is None or abs(cur - float(expected)) > 0.05:
                issues.append({
                    "location": "页面设置",
                    "element": names.get(key, key),
                    "current": _fmt_value(key, cur),
                    "expected": _fmt_value(key, float(expected)),
                    "suggestion": f"请在页面设置中将{names.get(key, key)}调整为{_fmt_value(key, float(expected))}。",
                    "source": "rule", "severity": rule.get("severity", "error"),
                    "paragraph_index": None,
                    "fix_hint": {"type": "page", "field": key, "value": float(expected)},
                })
        return issues

    # ---- 段落 ----
    def _check_paragraph(self, p, idx: int, rule: Dict, location: str) -> List[Dict]:
        issues = []
        checks = rule.get("checks", {})
        # 取该段落第一个非空 run 作为代表（公文通常同段字体一致）
        run = next((r for r in p.runs if r.text.strip()), None)
        pf = p.paragraph_format

        def add(element, kind, cur, expected, suggestion):
            issues.append({
                "location": location,
                "element": element,
                "current": _fmt_value(kind, cur),
                "expected": _fmt_value(kind, expected),
                "suggestion": suggestion,
                "source": "rule", "severity": rule.get("severity", "error"),
                "paragraph_index": idx,
                "fix_hint": {"type": "paragraph", "field": kind, "value": expected},
            })

        if "font_name" in checks and run is not None:
            cur = _run_font_name(run)
            if cur != checks["font_name"]:
                add("字体", "font_name", cur, checks["font_name"],
                    f"应将字体设置为「{checks['font_name']}」。")

        if "font_size_pt" in checks and run is not None:
            cur = run.font.size.pt if run.font.size else None
            if cur is None or abs(cur - float(checks["font_size_pt"])) > 0.5:
                add("字号", "font_size_pt", cur, float(checks["font_size_pt"]),
                    f"应将字号调整为{_fmt_value('font_size_pt', float(checks['font_size_pt']))}。")

        if "bold" in checks and run is not None:
            cur = bool(run.bold)
            if cur != bool(checks["bold"]):
                add("加粗", "bold", cur, bool(checks["bold"]),
                    "应设置为加粗。" if checks["bold"] else "不应加粗。")

        if "alignment" in checks:
            expected_align = _ALIGN_MAP.get(checks["alignment"])
            cur_align = p.alignment
            # 未显式设置对齐时 Word 默认左对齐
            cur_key = _ALIGN_NAME.get(cur_align, "left")
            if cur_key != checks["alignment"]:
                add("对齐方式", "alignment", cur_key, checks["alignment"],
                    f"应设置为{_ALIGN_CN.get(checks['alignment'], checks['alignment'])}。")

        if "line_spacing_pt" in checks:
            cur = pf.line_spacing.pt if hasattr(pf.line_spacing, 'pt') and pf.line_spacing else None
            expected = float(checks["line_spacing_pt"])
            if cur is None or abs(cur - expected) > 0.5:
                add("行距", "line_spacing_pt", cur, expected,
                    f"应将行距调整为固定值{_fmt_value('line_spacing_pt', expected)}。")

        if "first_line_indent_chars" in checks:
            expected_chars = float(checks["first_line_indent_chars"])
            # docx 中缩进以长度存储；按正文字号估算字符宽度
            font_pt = None
            if run is not None and run.font.size:
                font_pt = run.font.size.pt
            if font_pt is None and "font_size_pt" in checks:
                font_pt = float(checks["font_size_pt"])
            cur_chars = None
            if pf.first_line_indent is not None and font_pt:
                # 1 字符 ≈ 字号磅值
                cur_chars = round(pf.first_line_indent.pt / font_pt, 1)
            if cur_chars is None or abs(cur_chars - expected_chars) > 0.3:
                add("首行缩进", "first_line_indent_chars", cur_chars if cur_chars is not None else 0,
                    expected_chars,
                    f"应设置首行缩进{expected_chars:g}字符。")

        for key, label in (("space_before_pt", "段前间距"), ("space_after_pt", "段后间距")):
            if key in checks:
                cur_val = getattr(pf, "space_before" if "before" in key else "space_after")
                cur_pt = cur_val.pt if cur_val else None
                expected = float(checks[key])
                if cur_pt is None or abs(cur_pt - expected) > 0.5:
                    add(label, key, cur_pt, expected,
                        f"应将{label}调整为{_fmt_value(key, expected)}。")
        return issues

    # ---- 全文通用 ----
    def _check_general(self, paragraphs, rule: Dict) -> List[Dict]:
        issues = []
        checks = rule.get("checks", {})

        if checks.get("no_extra_blank_lines"):
            blank_run = 0
            for idx, p in enumerate(paragraphs):
                if not p.text.strip():
                    blank_run += 1
                    if blank_run >= 2:
                        issues.append({
                            "location": f"第{idx + 1}段附近",
                            "element": "多余空行",
                            "current": f"连续 {blank_run} 个空行",
                            "expected": "不允许多余空行",
                            "suggestion": "删除多余的空行，段落间距应通过段前/段后设置实现，不要用回车空行。",
                            "source": "rule", "severity": rule.get("severity", "warning"),
                            "paragraph_index": idx,
                            "fix_hint": {"type": "delete_paragraph"},
                        })
                else:
                    blank_run = 0

        if checks.get("no_trailing_spaces"):
            for idx, p in enumerate(paragraphs):
                if p.text != p.text.rstrip() and p.text.strip():
                    issues.append({
                        "location": f"第{idx + 1}段",
                        "element": "行尾空格",
                        "current": "段末存在多余空格",
                        "expected": "段末不应有空格",
                        "suggestion": "删除段末多余空格。",
                        "source": "rule", "severity": rule.get("severity", "warning"),
                        "paragraph_index": idx,
                        "fix_hint": {"type": "trim_text"},
                    })
        return issues

    # ================= 纯文本（txt/md） =================
    def _check_plain_text(self, text: str, filename: str, rules: List[Dict],
                          use_ai: bool) -> Dict:
        issues = [{
            "location": "全文",
            "element": "文件类型",
            "current": "纯文本文件",
            "expected": "Word (.docx) 文件",
            "suggestion": "纯文本不包含字体/字号等排版属性，仅进行内容层面检查。",
            "source": "rule", "severity": "warning",
            "paragraph_index": None, "fix_hint": None
        }]
        ai_used = False
        if use_ai and self.llm is not None:
            ai_issues = self._ai_check(text[:FORMAT_CHECK_AI_BUDGET], rules)
            if ai_issues is not None:
                issues.extend(ai_issues)
                ai_used = True
        return {"filename": filename, "file_type": "txt", "issues": issues, "ai_used": ai_used}

    # ================= AI 辅助判断 =================
    def _ai_check(self, text: str, rules: List[Dict]) -> Optional[List[Dict]]:
        """让 AI 判断规则难以表达的格式/规范问题。返回 None 表示 AI 不可用。"""
        if not text.strip():
            return []
        rule_brief = "；".join(
            f"{r['name']}（{r['target']}）" for r in rules[:15]
        ) or "（暂无已配置规则）"
        prompt = f"""你是公文格式审核助手。以下是一份公文的纯文本内容（排版属性已另行检查，无需关注字体字号）。
请只检查这些规则难以判断的问题：
1. 落款（发文机关署名）和成文日期是否缺失或明显位置不当；
2. 标题与正文结构是否明显混乱；
3. 是否存在明显的重复段落、残留占位符（如"××"、"XXX"、"待补充"）；
4. 文号格式是否明显不规范（若有文号）。

已配置的格式规则供参考：{rule_brief}

公文内容：
{text}

请严格按以下 JSON 数组格式输出（不要输出其他任何内容），没有问题则输出 []：
[{{"location": "位置描述", "element": "问题项", "current": "当前情况", "expected": "要求", "suggestion": "修改建议"}}]"""
        try:
            reply = self.llm._call_vllm(
                [{"role": "system", "content": "你是严谨的公文格式审核助手，只输出 JSON。"},
                 {"role": "user", "content": prompt}],
                temperature=0.1, max_tokens=1024
            )
            if not reply or ("调用" in reply and "失败" in reply):
                return None
            m = re.search(r'\[.*\]', reply, re.S)
            if not m:
                return []
            items = _json.loads(m.group(0))
            result = []
            for it in items[:20]:
                result.append({
                    "location": str(it.get("location", "全文")),
                    "element": str(it.get("element", "格式问题")),
                    "current": str(it.get("current", "")),
                    "expected": str(it.get("expected", "")),
                    "suggestion": str(it.get("suggestion", "")),
                    "source": "ai",
                    "severity": "warning",
                    "paragraph_index": None,
                    "fix_hint": None,
                })
            return result
        except Exception as e:
            print(f"[FormatCheck] AI 辅助判断失败: {e}")
            return None

    # ================= 自动修正 =================
    def apply_fixes(self, file_path: Path, issues: List[Dict],
                    accepted_indices: Optional[List[int]] = None,
                    out_path: Optional[Path] = None) -> List[str]:
        """按 issue 中的 fix_hint 自动修正 docx。

        :param file_path: 源 docx 路径
        :param issues: 校验返回的 issue 列表（顺序与前端展示一致）
        :param accepted_indices: 要应用的 issue 下标；传 None 表示全部应用
        :param out_path: 修正稿输出路径；传 None 则在源文件同名目录生成 *_fixed.docx
        :return: 修正后的段落文本列表（供审阅模式右栏预览）
        """
        file_path = Path(file_path)
        if out_path is None:
            out_path = file_path.with_name(file_path.stem + "_fixed.docx")
        out_path = Path(out_path)

        if accepted_indices is None:
            accepted_indices = list(range(len(issues)))

        doc = DocxDocument(str(file_path))

        # 先应用非删除类修正；删除段落最后按索引从大到小执行，避免索引位移
        deletes = []
        for idx in accepted_indices:
            if idx < 0 or idx >= len(issues):
                continue
            issue = issues[idx] or {}
            hint = issue.get("fix_hint") or {}
            htype = hint.get("type")
            p_idx = issue.get("paragraph_index")
            try:
                if htype == "page":
                    self._apply_page_fix(doc, hint.get("field"), hint.get("value"))
                elif htype == "paragraph":
                    if p_idx is None or p_idx >= len(doc.paragraphs):
                        continue
                    self._apply_paragraph_fix(doc.paragraphs[p_idx],
                                              hint.get("field"), hint.get("value"))
                elif htype == "trim_text":
                    if p_idx is None or p_idx >= len(doc.paragraphs):
                        continue
                    for r in doc.paragraphs[p_idx].runs:
                        if r.text != r.text.rstrip():
                            r.text = r.text.rstrip()
                elif htype == "delete_paragraph":
                    if p_idx is not None:
                        deletes.append(p_idx)
            except Exception as e:
                print(f"[FormatCheck] 修正第{idx}条失败（已跳过）: {e}")

        for p_idx in sorted(set(deletes), reverse=True):
            try:
                if p_idx < len(doc.paragraphs):
                    p = doc.paragraphs[p_idx]._p
                    p.getparent().remove(p)
            except Exception as e:
                print(f"[FormatCheck] 删除第{p_idx + 1}段失败（已跳过）: {e}")

        doc.save(str(out_path))

        # 返回修正后的段落文本供预览
        return [p.text for p in DocxDocument(str(out_path)).paragraphs]

    # ---- 修正辅助 ----
    @staticmethod
    def _set_run_font(run, font_name=None, size_pt=None, bold=None):
        if font_name:
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
            rfonts.set(qn("w:eastAsia"), font_name)
        if size_pt is not None:
            run.font.size = Pt(float(size_pt))
        if bold is not None:
            run.font.bold = bool(bold)

    def _apply_paragraph_fix(self, para, field, value):
        pf = para.paragraph_format
        if field in ("font_name", "font_size_pt", "bold"):
            if not para.runs:
                para.add_run("")
            for r in para.runs:
                if field == "font_name":
                    self._set_run_font(r, font_name=str(value))
                elif field == "font_size_pt":
                    self._set_run_font(r, size_pt=float(value))
                else:
                    self._set_run_font(r, bold=bool(value))
        elif field == "alignment":
            align = _ALIGN_MAP.get(str(value).strip().lower())
            if align is not None:
                pf.alignment = align
        elif field == "line_spacing_pt":
            pf.line_spacing = Pt(float(value))
        elif field == "first_line_indent_chars":
            chars = float(value)
            pf.first_line_indent = Pt(12 * chars)  # 兜底近似
            ppr = para._p.get_or_add_pPr()
            ind = ppr.find(qn("w:ind"))
            if ind is None:
                ind = ppr.makeelement(qn("w:ind"), {})
                ppr.append(ind)
            ind.set(qn("w:firstLineChars"), str(int(chars * 100)))
        elif field == "space_before_pt":
            pf.space_before = Pt(float(value))
        elif field == "space_after_pt":
            pf.space_after = Pt(float(value))

    @staticmethod
    def _apply_page_fix(doc, field, value):
        """页面级修正：页边距/页面尺寸（*_cm 字段，单位厘米）。"""
        attr = re.sub(r'_cm$', '', str(field or ''))
        if not attr:
            return
        for sec in doc.sections:
            try:
                if hasattr(sec, attr):
                    setattr(sec, attr, Cm(float(value)))
            except Exception:
                pass