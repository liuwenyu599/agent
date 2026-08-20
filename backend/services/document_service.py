
# -*- coding: utf-8 -*-
"""文档入库服务（覆盖 backend/services/document_service.py）

本次改动（网页链接导入知识库）：
- 抽出 _persist_document()：分块、建 Chunk、入库、向量化 的公共流程；
- process_upload() 行为不变，改走 _persist_document()；
- 新增 process_web_document()：网页抓取结果直接入库，与文件上传走完全相同的
  分块 → Embedding → FAISS 流程，保证 RAG 行为一致（需求三）。
- 网页文档元数据（source_url / source_type=web / source_name / publish_time）
  存入 doc_metadata，并支持按 source_url 查重。
"""
import re
import os
import uuid
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from backend.database.models import Document as DBDocument, Chunk
from backend.rag.chunker.legal_chunker import LegalChunker
from backend.rag.chunker.official_chunker import OfficialChunker

class DocumentService:
    def __init__(self, embedder=None):
        self.embedder = embedder
        self.legal_chunker = LegalChunker()
        self.official_chunker = OfficialChunker()

    def _detect_doc_type(self, text: str) -> str:
        """更准确的文档类型自动识别"""
        sample = text[:10000]

        # 法规：必须有"第X条"格式 + 法规关键词
        if re.search(r'第[一二三四五六七八九十百千零\d]+条[、\.\s]', sample):
            if any(k in sample for k in ["办法", "规定", "条例", "细则", "章程", "通则", "规则"]):
                return "法规"

        # 执法文书
        if any(k in sample for k in ["笔录", "告知书", "决定书", "执法", "行政处罚", "责令改正"]):
            return "执法文书"

        # 公文：常见的公文类型关键词
        official_keywords = ["通知", "通报", "报告", "请示", "批复", "函", "纪要", "决定", "意见",
                             "工作总结", "工作计划", "实施方案", "工作要点", "年度总结", "述职报告",
                             "会议纪要", "会议记录", "领导讲话", "调研报告", "情况汇报"]
        if any(k in sample for k in official_keywords):
            return "公文"

        # 默认公文（党政机关上传的文件大概率是公文）
        return "公文"

    # ========== 公共入库流程（文件 / 网页共用） ==========

    def _persist_document(self, db: Session, kb_id: str, title: str, text: str,
                          user_id: str, user_role: str,
                          doc_type: Optional[str] = None,
                          file_path: Optional[str] = None,
                          file_size: int = 0,
                          doc_metadata: Optional[dict] = None,
                          department: Optional[str] = None,
                          doc_number: Optional[str] = None) -> Dict:
        """分块 → 建 Document/Chunk → 向量化。文件上传和网页导入共用。"""
        doc_id = str(uuid.uuid4())[:8]

        # 自动识别文档类型（如果调用方没指定）
        auto_doc_type = self._detect_doc_type(text)
        final_doc_type = doc_type if doc_type else auto_doc_type

        # 根据文档类型选择切片器
        if final_doc_type == "法规":
            chunks = self.legal_chunker.chunk(text, doc_id)
        else:
            chunks = self.official_chunker.chunk(text, doc_id)

        # 确定初始状态：管理员/知识管理员直接发布；个人库直接发布；公共库待审核
        if user_role in ["knowledge_admin", "developer"]:
            initial_status = "published"
        else:
            from backend.database.models import KnowledgeBase as KBModel
            kb = db.query(KBModel).filter(KBModel.id == kb_id).first()
            if kb and kb.kb_type == "personal" and kb.owner_id == user_id:
                initial_status = "published"
            else:
                initial_status = "pending"

        doc = DBDocument(
            id=doc_id,
            kb_id=kb_id,
            title=title,
            doc_type=final_doc_type,
            department=department,
            doc_number=doc_number,
            file_path=file_path,
            file_size=file_size,
            status=initial_status,
            uploaded_by=user_id,
            created_by=user_id,
            content=text[:500000],
            doc_metadata=doc_metadata or {}
        )
        db.add(doc)

        # 创建 Chunk 记录，并确保 chunk_id 一致
        chunk_records = []
        for i, chunk_data in enumerate(chunks):
            chunk_id = f"{doc_id}_{i}"  # 统一使用 doc_id_index 格式
            chunk = Chunk(
                id=chunk_id,  # 数据库 ID 和 FAISS 一致
                doc_id=doc_id,
                chunk_index=chunk_data["chunk_index"],
                chunk_type=chunk_data["chunk_type"],
                title=chunk_data.get("title", ""),
                content=chunk_data["content"],
                char_count=chunk_data["metadata"].get("char_count", 0),
                word_count=chunk_data["metadata"].get("word_count", 0),
                chunk_metadata=chunk_data["metadata"]
            )
            db.add(chunk)
            chunk_records.append({
                "chunk_id": chunk_id,
                "doc_id": doc_id,
                "kb_id": kb_id,
                "content": chunk_data["content"],
                "title": chunk_data.get("title", "")
            })

        db.commit()

        # 向量化（仅对已发布文档）
        if initial_status == "published" and self.embedder and chunk_records:
            texts = [c["content"] for c in chunk_records]
            try:
                vectors = self.embedder.encode(texts)
                from backend.infrastructure.vectorstore.faiss_store import FAISSVectorStore
                vector_store = FAISSVectorStore()

                faiss_chunks = []
                for i, chunk_rec in enumerate(chunk_records):
                    faiss_chunks.append({
                        "chunk_id": chunk_rec["chunk_id"],
                        "doc_id": chunk_rec["doc_id"],
                        "kb_id": chunk_rec["kb_id"],
                        "embedding": vectors[i],
                        "content": chunk_rec["content"],
                        "title": chunk_rec["title"]
                    })
                vector_store.add_chunks(faiss_chunks)
                print(f"[FAISS] 已索引 {len(faiss_chunks)} 个 chunk")
            except Exception as e:
                print(f"向量化失败: {e}")

        return {
            "doc_id": doc_id,
            "chunks": len(chunks),
            "status": initial_status,
            "message": "已直接发布" if initial_status == "published" else "已提交审核，等待管理员审批"
        }

    # ========== 网页入库 ==========

    def process_web_document(self, db: Session, kb_id: str, user_id: str, user_role: str,
                             url: str, fetched: Dict,
                             title: Optional[str] = None,
                             source_name: Optional[str] = None,
                             publish_time: Optional[str] = None) -> Dict:
        """把已抓取成功的网页内容写入知识库。

        fetched: web_fetcher.fetch_webpage 的返回（ok=True）。
        title/source_name/publish_time：Excel 元数据优先，网页解析结果兜底。
        """
        text = (fetched.get("content") or "").strip()
        final_title = (title or fetched.get("title") or url).strip()[:500]

        # 把原始链接写进正文头部，检索时模型也能看到出处
        header = f"【来源网页】{url}"
        src = source_name or fetched.get("source_name") or ""
        if src:
            header += f"（{src}）"
        full_text = header + "\n\n" + text

        meta = {
            "source_type": "web",
            "source_url": url,
            "source_name": src,
            "publish_time": publish_time or fetched.get("publish_time") or "",
        }

        result = self._persist_document(
            db=db, kb_id=kb_id, title=final_title, text=full_text,
            user_id=user_id, user_role=user_role,
            doc_type="网页", file_path=None, file_size=len(text),
            doc_metadata=meta,
        )
        result["title"] = final_title
        result["source_url"] = url
        return result

    def find_by_source_url(self, db: Session, source_url: str) -> Optional[DBDocument]:
        """按 source_url 查重（需求：同一 URL 只保留一份，archived 的也算已存在）。"""
        try:
            return db.query(DBDocument).filter(
                DBDocument.doc_metadata["source_url"].as_string() == source_url
            ).first()
        except Exception:
            # 某些数据库方言不支持 JSON 路径查询，退化为内存过滤
            docs = db.query(DBDocument).filter(DBDocument.doc_metadata.isnot(None)).all()
            for d in docs:
                if (d.doc_metadata or {}).get("source_url") == source_url:
                    return d
            return None

    # ========== 文件上传（原逻辑，改走公共流程） ==========

    async def process_upload(self, file, kb_id: str, user_id: str, user_role: str, db: Session, title=None, doc_type=None, department=None, doc_number=None) -> Dict:
        doc_id = str(uuid.uuid4())[:8]

        # 保存临时文件
        tmp_dir = Path("/tmp/judicial-upload")
        tmp_dir.mkdir(exist_ok=True)
        tmp_path = tmp_dir / f"{doc_id}_{file.filename}"

        content = await file.read()
        with open(tmp_path, "wb") as f:
            f.write(content)

        # 解析文档
        text = self._extract_text(tmp_path, file.filename)

        result = self._persist_document(
            db=db, kb_id=kb_id,
            title=title if title else file.filename.replace(".docx", "").replace(".txt", "").replace(".pdf", "").replace(".md", ""),
            text=text, user_id=user_id, user_role=user_role,
            doc_type=doc_type, file_path=str(tmp_path), file_size=len(content),
            doc_metadata={"source_type": "file", "original_filename": file.filename},
            department=department, doc_number=doc_number,
        )

        # 清理临时文件
        os.remove(tmp_path)
        return result

    def review_document(self, doc_id: str, action: str, comment: str, reviewer_id: str, db: Session) -> Dict:
        """审核文档"""
        doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
        if not doc:
            return {"error": "Document not found"}

        if action in ["approve", "approved"]:
            doc.status = "published"
            doc.reviewed_by = reviewer_id
            doc.reviewed_at = datetime.utcnow()
            doc.review_comment = comment or "审核通过"

            # 已发布文档进行向量化
            if self.embedder:
                chunks = db.query(Chunk).filter(Chunk.doc_id == doc_id).all()
                if chunks:
                    texts = [c.content for c in chunks]
                    try:
                        vectors = self.embedder.encode(texts)
                        from backend.infrastructure.vectorstore.faiss_store import FAISSVectorStore
                        vector_store = FAISSVectorStore()

                        faiss_chunks = []
                        for i, chunk in enumerate(chunks):
                            faiss_chunks.append({
                                "chunk_id": chunk.id,
                                "doc_id": doc_id,
                                "kb_id": doc.kb_id,
                                "embedding": vectors[i],
                                "content": chunk.content,
                                "title": chunk.title or ""
                            })
                        vector_store.add_chunks(faiss_chunks)
                        print(f"[FAISS] 已索引 {len(faiss_chunks)} 个 chunk")
                    except Exception as e:
                        print(f"向量化失败: {e}")

        elif action in ["reject", "rejected"]:
            doc.status = "rejected"
            doc.reviewed_by = reviewer_id
            doc.reviewed_at = datetime.utcnow()
            doc.review_comment = comment or "审核未通过"

        db.commit()

        return {
            "doc_id": doc_id,
            "status": doc.status,
            "message": "审核完成"
        }

    def _extract_text(self, file_path: Path, filename: str) -> str:
        suffix = filename.lower().split('.')[-1]

        if suffix == 'docx':
            doc = DocxDocument(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        elif suffix == 'doc':
            try:
                import subprocess
                result = subprocess.run(['antiword', str(file_path)], capture_output=True, text=True)
                if result.returncode == 0:
                    return result.stdout
            except:
                pass
            try:
                import docx2txt
                return docx2txt.process(str(file_path))
            except:
                pass
            return "[.doc 格式解析失败，请转换为 .docx 后上传]"
        elif suffix == 'pdf':
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(str(file_path))
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except Exception as e:
                print(f"PDF 解析失败: {e}")
                return "[PDF 解析失败]"
        elif suffix in ['txt', 'md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except:
                return ""

