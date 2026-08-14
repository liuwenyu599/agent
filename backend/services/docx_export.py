import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


def _set_font(run, name, size, bold=False, color=None):
    """设置字体"""
    run.font.size = size
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn('w:eastAsia'), name)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def _set_paragraph_format(p, alignment=None, first_line_indent=None, 
                          line_spacing=None, space_before=None, space_after=None):
    """设置段落格式"""
    if alignment is not None:
        p.alignment = alignment
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after


def _add_red_header(doc, title=""):
    """添加公文红头（发文机关标识）"""
    p = doc.add_paragraph()
    _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    run = p.add_run(title if title else "××司法局文件")
    _set_font(run, '方正小标宋简体', Pt(22), color=RGBColor(0xFF, 0x00, 0x00))

    # 添加红色分割线
    p = doc.add_paragraph()
    _set_paragraph_format(p, space_after=Pt(12))
    # 通过下划线模拟分割线
    run = p.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    _set_font(run, '宋体', Pt(12), color=RGBColor(0xFF, 0x00, 0x00))


def _add_doc_number(doc, number=""):
    """添加发文字号"""
    p = doc.add_paragraph()
    _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    run = p.add_run(number if number else "××司发〔2026〕××号")
    _set_font(run, '仿宋', Pt(14))


def _add_recipient(doc, recipient=""):
    """添加主送机关"""
    if not recipient:
        return
    p = doc.add_paragraph()
    _set_paragraph_format(p, space_after=Pt(6))
    run = p.add_run(recipient)
    _set_font(run, '仿宋', Pt(16))


def _parse_inline_formatting(p, text):
    """解析行内格式：粗体、斜体、下划线"""
    # 先处理粗体 **text**
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            _set_font(run, '仿宋', Pt(16))
        else:
            # 处理斜体 *text*
            sub_parts = re.split(r'((?<!\*)\*(?!\*).+?\*(?!\*))', part)
            for sub in sub_parts:
                if not sub:
                    continue
                if sub.startswith('*') and sub.endswith('*') and len(sub) > 2:
                    run = p.add_run(sub[1:-1])
                    run.italic = True
                    _set_font(run, '仿宋', Pt(16))
                else:
                    run = p.add_run(sub)
                    _set_font(run, '仿宋', Pt(16))


def _add_table(doc, lines, start_idx):
    """解析并添加 Markdown 表格"""
    # 收集表格行
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1

    if len(table_lines) < 2:
        return i

    # 解析表头
    header_cells = [c.strip() for c in table_lines[0].split('|')[1:-1]]
    # 跳过分隔行（包含 --- 的那行）
    data_rows = []
    for line in table_lines[2:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            data_rows.append(cells)

    if not header_cells:
        return i

    # 创建表格
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_cells))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 设置表头
    hdr_cells = table.rows[0].cells
    for j, cell_text in enumerate(header_cells):
        if j < len(hdr_cells):
            hdr_cells[j].text = cell_text
            for paragraph in hdr_cells[j].paragraphs:
                _set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                for run in paragraph.runs:
                    _set_font(run, '黑体', Pt(12), bold=True)

    # 设置数据行
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        for j, cell_text in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].text = cell_text
                for paragraph in row_cells[j].paragraphs:
                    _set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    for run in paragraph.runs:
                        _set_font(run, '仿宋', Pt(12))

    return i


def _add_list_item(doc, line, level=0):
    """添加列表项"""
    p = doc.add_paragraph()
    indent = Cm(0.74 * level)  # 每级缩进 0.74cm
    _set_paragraph_format(
        p, 
        first_line_indent=Cm(0.74) if level == 0 else None,
        line_spacing=Pt(28),
        space_after=Pt(3)
    )

    # 去除列表标记
    content = re.sub(r'^\s*([-*]|\d+\.)\s+', '', line)

    # 添加缩进
    if level > 0:
        p.paragraph_format.left_indent = indent

    # 列表标记符号
    if line.strip().startswith('-') or line.strip().startswith('*'):
        prefix = '• '
    elif re.match(r'^\s*\d+\.', line):
        prefix = re.match(r'^\s*(\d+\.)', line).group(1) + ' '
    else:
        prefix = ''

    run = p.add_run(prefix)
    _set_font(run, '仿宋', Pt(16))
    _parse_inline_formatting(p, content)

    return p


def _add_blockquote(doc, lines, start_idx):
    """添加引用块"""
    p = doc.add_paragraph()
    _set_paragraph_format(
        p,
        first_line_indent=Cm(0.74),
        line_spacing=Pt(28),
        space_after=Pt(6)
    )

    # 收集引用内容
    quote_text = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('>'):
        text = lines[i].strip()[1:].strip()
        quote_text.append(text)
        i += 1

    full_text = ' '.join(quote_text)
    run = p.add_run(full_text)
    _set_font(run, '楷体', Pt(15), color=RGBColor(0x66, 0x66, 0x66))

    return i


def _add_code_block(doc, lines, start_idx):
    """添加代码块"""
    p = doc.add_paragraph()
    _set_paragraph_format(p, space_after=Pt(6))

    code_lines = []
    i = start_idx
    while i < len(lines) and not lines[i].strip().startswith('```'):
        code_lines.append(lines[i])
        i += 1

    # 跳过结束标记
    if i < len(lines) and lines[i].strip().startswith('```'):
        i += 1

    full_code = '\n'.join(code_lines)
    run = p.add_run(full_code)
    _set_font(run, 'Courier New', Pt(12), color=RGBColor(0x33, 0x33, 0x33))

    # 设置代码块背景色（通过 shading）
    shading_elm = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
    p._element.get_or_add_pPr().append(shading_elm)

    return i


def _add_horizontal_rule(doc):
    """添加水平分割线"""
    p = doc.add_paragraph()
    _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(6))
    run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    _set_font(run, '宋体', Pt(12), color=RGBColor(0xCC, 0xCC, 0xCC))


def _add_page_setup(doc):
    """设置页面格式（A4，标准页边距）"""
    for section in doc.sections:
        section.page_width = Cm(21.0)   # A4 宽度
        section.page_height = Cm(29.7)  # A4 高度
        section.top_margin = Cm(3.7)    # 上 37mm
        section.bottom_margin = Cm(3.5)  # 下 35mm
        section.left_margin = Cm(2.8)   # 左 28mm
        section.right_margin = Cm(2.6)  # 右 26mm


def _add_header_footer(doc, title=""):
    """添加页眉页脚"""
    for section in doc.sections:
        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = title if title else "白云司法智能写作助手"
        _set_paragraph_format(hp, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for run in hp.runs:
            _set_font(run, '宋体', Pt(10), color=RGBColor(0x80, 0x80, 0x80))

        # 页脚 - 页码
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        _set_paragraph_format(fp, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # 添加页码字段
        run = fp.add_run('— ')
        _set_font(run, '宋体', Pt(10))

        # 添加 PAGE 字段
        fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
        run1 = fp.add_run()
        run1._element.append(fldChar1)

        instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w')))
        run2 = fp.add_run()
        run2._element.append(instrText)

        fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
        run3 = fp.add_run()
        run3._element.append(fldChar2)

        run = fp.add_run(' —')
        _set_font(run, '宋体', Pt(10))


def _add_signature(doc, signature_text=""):
    """添加落款（发文机关署名、成文日期）"""
    if not signature_text:
        return

    p = doc.add_paragraph()
    _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(24), space_after=Pt(6))
    run = p.add_run(signature_text)
    _set_font(run, '仿宋', Pt(16))


def _add_date(doc, date_text=""):
    """添加成文日期"""
    if not date_text:
        return

    p = doc.add_paragraph()
    _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(12))
    run = p.add_run(date_text)
    _set_font(run, '仿宋', Pt(16))


def markdown_to_docx(text: str, title: str = "公文", 
                     doc_number: str = "",
                     recipient: str = "",
                     signature: str = "",
                     date_text: str = "",
                     use_red_header: bool = False) -> BytesIO:
    """
    将 Markdown 文本转换为符合党政机关公文格式的 Word 文档

    Args:
        text: Markdown 格式的文本内容
        title: 文档标题
        doc_number: 发文字号（如：××司发〔2026〕1号）
        recipient: 主送机关
        signature: 落款单位
        date_text: 成文日期
        use_red_header: 是否使用红头格式

    Returns:
        BytesIO: 包含 Word 文档的字节流
    """
    doc = Document()

    # 设置页面
    _add_page_setup(doc)

    # 添加页眉页脚
    _add_header_footer(doc, title)

    # 红头（可选）
    if use_red_header:
        _add_red_header(doc, title)
        _add_doc_number(doc, doc_number)
    else:
        # 普通标题
        if title and title != "公文":
            p = doc.add_paragraph()
            _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
            run = p.add_run(title)
            _set_font(run, '方正小标宋简体', Pt(22), bold=True)

    # 主送机关
    if recipient:
        _add_recipient(doc, recipient)

    # 解析 Markdown 内容
    lines = text.split('\n')
    i = 0
    in_code_block = False
    in_list = False
    list_level = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                i += 1
                i = _add_code_block(doc, lines, i)
                in_code_block = False
                continue

        if in_code_block:
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 水平分割线
        if stripped == '---' or stripped == '***' or stripped == '___':
            _add_horizontal_rule(doc)
            i += 1
            continue

        # 表格
        if stripped.startswith('|'):
            i = _add_table(doc, lines, i)
            continue

        # 引用块
        if stripped.startswith('>'):
            i = _add_blockquote(doc, lines, i)
            continue

        # 标题
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
            run = p.add_run(stripped[2:].strip())
            _set_font(run, '方正小标宋简体', Pt(22), bold=True)
            i += 1
            continue

        if stripped.startswith('## '):
            p = doc.add_paragraph()
            _set_paragraph_format(p, space_before=Pt(12), space_after=Pt(6))
            run = p.add_run(stripped[3:].strip())
            _set_font(run, '黑体', Pt(16), bold=True)
            i += 1
            continue

        if stripped.startswith('### '):
            p = doc.add_paragraph()
            _set_paragraph_format(p, space_before=Pt(6), space_after=Pt(3))
            run = p.add_run(stripped[4:].strip())
            _set_font(run, '楷体', Pt(15), bold=True)
            i += 1
            continue

        # 列表项
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.+)$', line)
        if list_match:
            indent_len = len(list_match.group(1))
            level = indent_len // 2  # 每2个空格一级
            _add_list_item(doc, line, level)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _set_paragraph_format(
            p,
            first_line_indent=Cm(0.74),
            line_spacing=Pt(28),
            space_after=Pt(3)
        )
        _parse_inline_formatting(p, stripped)
        i += 1

    # 落款
    if signature:
        _add_signature(doc, signature)

    # 日期
    if date_text:
        _add_date(doc, date_text)

    # 保存
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_official_document(content: str, 
                                title: str = "",
                                doc_number: str = "",
                                recipient: str = "",
                                signature: str = "",
                                date_text: str = "") -> BytesIO:
    """
    生成标准党政机关公文格式 Word 文档

    这是 markdown_to_docx 的便捷封装，默认启用红头格式
    """
    return markdown_to_docx(
        text=content,
        title=title,
        doc_number=doc_number,
        recipient=recipient,
        signature=signature,
        date_text=date_text,
        use_red_header=True
    )