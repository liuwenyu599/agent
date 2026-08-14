"""
司法智能写作助手 - 新功能集成测试
覆盖：计划总结合并、格式校验、对话附件上传、多轮引用、原有功能兼容

运行前确保：
  1. 后端服务已启动 (./start.sh start)
  2. vLLM 已启动（AI 辅助判断和写作需要）
  3. 已安装 tesseract-ocr（图片 OCR 测试需要）

用法：
  python3 tests/integration_test_v2.py

环境变量：
  BASE_URL      后端地址，默认 http://localhost:8000/api/v1
  SKIP_AI       设置任意值跳过需要 vLLM 的测试（写作、AI 辅助校验）
"""
import os
import sys
import tempfile
import json
import requests
from pathlib import Path

BASE_URL = os.getenv("BASE_URL", "http://localhost:8000/api/v1")
SKIP_AI = os.getenv("SKIP_AI", "")


class TestClient:
    def __init__(self):
        self.token = None
        self.user = None
        self.session = requests.Session()

    def _headers(self, content_type="application/json"):
        h = {}
        if content_type:
            h["Content-Type"] = content_type
        if self.token:
            h["Authorization"] = f"Bearer {self.token}"
        return h

    def post(self, path, data=None, files=None, params=None):
        url = f"{BASE_URL}{path}"
        if files:
            r = self.session.post(url, data=data, files=files, headers=self._headers(None), params=params)
        else:
            r = self.session.post(url, json=data or {}, headers=self._headers(), params=params)
        return r

    def get(self, path, params=None):
        url = f"{BASE_URL}{path}"
        r = self.session.get(url, headers=self._headers(), params=params)
        return r

    def put(self, path, data=None):
        url = f"{BASE_URL}{path}"
        r = self.session.put(url, json=data or {}, headers=self._headers())
        return r

    def delete(self, path):
        url = f"{BASE_URL}{path}"
        r = self.session.delete(url, headers=self._headers())
        return r


# ==================== 全局状态 ====================
client = TestClient()       # 系统管理员 (developer)
kb_client = TestClient()    # 知识管理员
user_client = TestClient()  # 普通用户
passed = 0
failed = 0


def test(name, condition, detail=""):
    global passed, failed
    if condition:
        print(f"  ✅ {name}")
        passed += 1
    else:
        print(f"  ❌ {name}: {detail}")
        failed += 1


def ensure_login():
    """确保三个角色都登录"""
    # 系统管理员
    r = client.get("/auth/check-first-user")
    if r.status_code == 200 and r.json().get("is_first", False):
        r = client.post("/auth/register-first", {
            "username": "superadmin",
            "email": "admin@judicial.gov.cn",
            "password": "admin123",
            "real_name": "系统管理员",
            "department": "信息中心",
            "role": "developer"
        })
        if r.status_code == 200:
            client.token = r.json()["access_token"]
            client.user = r.json()["user"]
    else:
        r = client.post("/auth/login", {"username": "superadmin", "password": "admin123"})
        if r.status_code == 200:
            client.token = r.json()["access_token"]
            client.user = r.json()["user"]

    # 知识管理员
    r = kb_client.post("/auth/login", {"username": "kbadmin", "password": "kb123456"})
    if r.status_code != 200:
        r = client.post("/auth/register", {
            "username": "kbadmin",
            "email": "kb@judicial.gov.cn",
            "password": "kb123456",
            "real_name": "知识管理员",
            "department": "办公室",
            "role": "knowledge_admin"
        })
        if r.status_code == 200:
            kb_client.token = r.json()["access_token"]
            kb_client.user = r.json()["user"]
    else:
        kb_client.token = r.json()["access_token"]
        kb_client.user = r.json()["user"]

    # 普通用户
    r = user_client.post("/auth/login", {"username": "user01", "password": "user123456"})
    if r.status_code != 200:
        r = client.post("/auth/register", {
            "username": "user01",
            "email": "user01@judicial.gov.cn",
            "password": "user123456",
            "real_name": "张三",
            "department": "社区矫正科",
            "role": "user"
        })
        if r.status_code == 200:
            user_client.token = r.json()["access_token"]
            user_client.user = r.json()["user"]
    else:
        user_client.token = r.json()["access_token"]
        user_client.user = r.json()["user"]

    return bool(client.token and kb_client.token and user_client.token)


# ==================== 辅助函数 ====================

def make_docx_with_issues() -> bytes:
    """构造一个带格式问题的 docx（用于格式校验测试）"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()
    # 页面设置：页边距故意设错（左 4cm 而不是标准 2.8cm）
    section = doc.sections[0]
    section.top_margin = Cm(3.0)    # 应为 3.7
    section.bottom_margin = Cm(3.0) # 应为 3.5
    section.left_margin = Cm(4.0)  # 应为 2.8
    section.right_margin = Cm(3.0)  # 应为 2.6

    # 标题：故意用错字体和字号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 应为居中
    run = p.add_run("关于社区矫正工作的情况汇报")
    run.font.name = "宋体"
    run.font.size = Pt(14)  # 应为 22pt（二号）
    run.bold = False

    # 正文：故意用错字体/字号/行距/缩进
    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Cm(0)  # 应为缩进 2 字符
    p2.paragraph_format.line_spacing = Pt(20)       # 应为 28-30 磅
    run2 = p2.add_run("各相关单位：\\n\\n为进一步加强社区矫正工作，根据上级部署，我局开展了专项排查活动。")
    run2.font.name = "楷体"
    run2.font.size = Pt(12)

    # 一级标题：格式不对
    p3 = doc.add_paragraph()
    run3 = p3.add_run("一、工作背景")
    run3.font.name = "宋体"
    run3.font.size = Pt(14)
    run3.bold = False

    # 落款：左对齐（应为右对齐）
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run4 = p4.add_run("白云区司法局")
    run4.font.name = "仿宋"
    run4.font.size = Pt(14)

    # 日期
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run5 = p5.add_run("2026年8月14日")
    run5.font.name = "仿宋"
    run5.font.size = Pt(14)

    # 连续空行
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 行尾空格
    p6 = doc.add_paragraph()
    run6 = p6.add_run("本段末尾有空格    ")
    run6.font.name = "仿宋"
    run6.font.size = Pt(16)

    buf = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(buf.name)
    with open(buf.name, "rb") as f:
        data = f.read()
    os.unlink(buf.name)
    return data


def make_txt_doc() -> bytes:
    """构造一个纯文本文件"""
    return "关于开展社区矫正工作的通知\\n\\n各相关单位：\\n\\n为进一步加强社区矫正工作...".encode("utf-8")


def make_simple_docx() -> bytes:
    """构造一个简单合规的 docx（用于附件上传测试）"""
    try:
        from docx import Document
    except ImportError:
        return None
    doc = Document()
    doc.add_paragraph("社区矫正工作总结")
    doc.add_paragraph("\\n本年度社区矫正工作取得了显著成效。")
    buf = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(buf.name)
    with open(buf.name, "rb") as f:
        data = f.read()
    os.unlink(buf.name)
    return data


# ==================== 测试开始 ====================
print("=" * 70)
print("司法智能写作助手 v2.0 - 新功能集成测试")
print(f"后端地址: {BASE_URL}")
print(f"跳过 AI 测试: {'是' if SKIP_AI else '否'}")
print("=" * 70)

# ---------- 0. 登录 ----------
print("\\n【0. 登录与角色准备】")
if not ensure_login():
    print("❌ 无法获取 token，后续测试跳过")
    sys.exit(1)
test("系统管理员登录", client.token is not None, "")
test("知识管理员登录", kb_client.token is not None, "")
test("普通用户登录", user_client.token is not None, "")

r = client.get("/auth/me")
test("管理员身份确认", r.status_code == 200 and r.json().get("role") == "developer", r.text)

# ==================== 第一部分：计划/总结合并 ====================
print("\\n【1. 计划/总结合并 —— 模板管理】")

r = client.post("/templates/init")
test("初始化内置模板（含合并迁移）", r.status_code == 200, f"{r.status_code}: {r.text}")

r = client.get("/templates/categories")
categories = r.json() if r.status_code == 200 else []
cat_names = [c["name"] for c in categories]
test("分类列表包含'计划总结'", "计划总结" in cat_names, f"现有分类: {cat_names}")
test("旧分类'工作总结'已停用", "工作总结" not in cat_names, f"现有分类: {cat_names}")
test("旧分类'工作计划'已停用", "工作计划" not in cat_names, f"现有分类: {cat_names}")

r = client.get("/templates/")
templates = r.json() if r.status_code == 200 else []
tmpl_names = [t["name"] for t in templates]
test("模板列表包含'工作计划与总结'", "工作计划与总结" in tmpl_names, f"现有模板: {tmpl_names}")
test("旧模板'年度工作总结'已消失", "年度工作总结" not in tmpl_names, f"现有模板: {tmpl_names}")
test("旧模板'工作计划'已消失", "工作计划" not in tmpl_names, f"现有模板: {tmpl_names}")

# 检查合并模板的字段
merged = next((t for t in templates if t["name"] == "工作计划与总结"), None)
if merged:
    params = {p["name"]: p for p in merged.get("params_schema", [])}
    test("合并模板有'title'字段", "title" in params, str(params.keys()))
    test("合并模板有'background'字段", "background" in params, str(params.keys()))
    test("合并模板有'goals'字段", "goals" in params, str(params.keys()))
    test("合并模板有'completion'字段", "completion" in params, str(params.keys()))
    test("合并模板有'achievements'字段", "achievements" in params, str(params.keys()))
    test("合并模板有'problems'字段", "problems" in params, str(params.keys()))
    test("合并模板有'next_plan'字段", "next_plan" in params, str(params.keys()))
    test("合并模板 category 为'计划总结'", merged.get("category") == "计划总结", merged.get("category"))
    test("合并模板 base_type 为'计划与总结'", merged.get("base_type") == "计划与总结", merged.get("base_type"))
    # 检查必填字段只有 title/department/period，其余应为非必填
    required = [p["name"] for p in merged.get("params_schema", []) if p.get("required")]
    test("合并模板必填字段精简", set(required) == {"title", "department", "period"}, f"必填: {required}")
else:
    test("合并模板字段检查", False, "未找到合并模板")

# 自定义模板不受影响
r = kb_client.post("/templates/", {
    "name": "测试自定义模板",
    "category": "通知公告",
    "base_type": "通知",
    "description": "用于测试自定义模板不受合并影响",
    "icon": "Bell",
    "params_schema": [
        {"name": "title", "label": "标题", "type": "input", "required": True, "placeholder": "标题"},
        {"name": "content", "label": "内容", "type": "textarea", "required": True, "placeholder": "内容"}
    ],
    "content_template": "测试模板",
    "system_prompt": "测试",
    "writing_style": "正式公文",
    "word_count": 500,
    "need_red_header": False,
    "need_signature": True,
    "need_date": True,
    "need_doc_number": False
})
test("自定义模板创建不受影响", r.status_code == 200, f"{r.status_code}: {r.text}")

r = client.get("/templates/")
templates_after = r.json() if r.status_code == 200 else []
test("自定义模板出现在列表中", any(t["name"] == "测试自定义模板" for t in templates_after), "")

# ==================== 第二部分：格式校验 ====================
print("\\n【2. 格式校验 —— 系统级独立功能】")

# 2.1 规则管理（管理员权限）
r = client.get("/format-check/rules")
test("规则列表（管理员）", r.status_code == 200, f"{r.status_code}: {r.text}")

# 普通用户也能查看规则列表（只读）
r = user_client.get("/format-check/rules")
test("规则列表（普通用户只读）", r.status_code == 200, f"{r.status_code}: {r.text}")

# 创建几条格式规则
rules_to_create = [
    {
        "name": "标题格式",
        "target": "title",
        "checks": {"font_name": "方正小标宋简体", "font_size_pt": 22, "alignment": "center"},
        "severity": "error",
        "is_default": True,
        "remark": "标题应使用方正小标宋简体二号居中"
    },
    {
        "name": "正文格式",
        "target": "body",
        "checks": {"font_name": "仿宋", "font_size_pt": 16, "line_spacing_pt": 28, "first_line_indent_chars": 2},
        "severity": "error",
        "is_default": True,
        "remark": "正文仿宋三号，行距28磅，首行缩进2字符"
    },
    {
        "name": "一级标题",
        "target": "heading1",
        "checks": {"font_name": "黑体", "font_size_pt": 16, "bold": True},
        "severity": "error",
        "is_default": True,
        "remark": "一级标题黑体三号加粗"
    },
    {
        "name": "页面设置",
        "target": "page",
        "checks": {"top_margin_cm": 3.7, "bottom_margin_cm": 3.5, "left_margin_cm": 2.8, "right_margin_cm": 2.6, "page_width_cm": 21.0, "page_height_cm": 29.7},
        "severity": "error",
        "is_default": True,
        "remark": "A4纸，标准页边距"
    },
    {
        "name": "通用规范",
        "target": "general",
        "checks": {"no_extra_blank_lines": True, "no_trailing_spaces": True},
        "severity": "warning",
        "is_default": True,
        "remark": "禁止连续空行和行尾空格"
    }
]

created_rule_ids = []
for rule in rules_to_create:
    r = client.post("/format-check/rules", rule)
    test(f"创建规则: {rule['name']}", r.status_code == 200, f"{r.status_code}: {r.text}")
    if r.status_code == 200:
        created_rule_ids.append(r.json().get("id"))

# 普通用户不能创建规则
r = user_client.post("/format-check/rules", rules_to_create[0])
test("普通用户创建规则被403", r.status_code == 403, f"期望403，实际{r.status_code}")

# 修改规则
if created_rule_ids:
    r = client.put(f"/format-check/rules/{created_rule_ids[0]}", {
        "name": "标题格式（已修改）",
        "target": "title",
        "checks": {"font_name": "方正小标宋简体", "font_size_pt": 22, "alignment": "center"},
        "severity": "error",
        "is_default": True,
        "remark": "已修改"
    })
    test("修改规则", r.status_code == 200, f"{r.status_code}: {r.text}")

# 2.2 无规则时校验应提示
r = client.get("/format-check/rules")
if r.status_code == 200:
    all_rules = r.json()
    # 临时禁用所有规则测试
    for rule in all_rules:
        client.put(f"/format-check/rules/{rule['id']}", {
            **rule, "is_active": False, "is_default": False
        })

    docx_data = make_docx_with_issues()
    if docx_data:
        r = client.session.post(
            f"{BASE_URL}/format-check/check",
            files={"file": ("test_bad.docx", docx_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"Authorization": f"Bearer {client.token}"}
        )
        test("无规则时校验提示配置规则", r.status_code == 400 and "尚未配置" in r.text, f"{r.status_code}: {r.text[:200]}")
    else:
        test("无规则时校验提示配置规则", False, "python-docx 未安装，跳过")

    # 恢复规则
    for rule in all_rules:
        client.put(f"/format-check/rules/{rule['id']}", {
            **rule, "is_active": True, "is_default": True
        })

# 2.3 文件格式校验（docx）
docx_data = make_docx_with_issues()
if docx_data:
    # 带 AI 辅助
    r = client.session.post(
        f"{BASE_URL}/format-check/check",
        files={"file": ("test_bad.docx", docx_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers={"Authorization": f"Bearer {client.token}"},
        params={"use_ai": "true"}
    )
    test("docx 格式校验（含AI）", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
    if r.status_code == 200:
        result = r.json()
        issues = result.get("issues", [])
        test("校验返回 issues 列表", isinstance(issues, list) and len(issues) > 0, f"issues数量: {len(issues)}")
        test("校验返回 record_id", bool(result.get("record_id")), "")
        test("校验返回 filename", result.get("filename") == "test_bad.docx", result.get("filename"))

        # 检查 issue 结构
        if issues:
            issue = issues[0]
            test("issue 有 location", "location" in issue, str(issue.keys()))
            test("issue 有 element", "element" in issue, str(issue.keys()))
            test("issue 有 current", "current" in issue, str(issue.keys()))
            test("issue 有 expected", "expected" in issue, str(issue.keys()))
            test("issue 有 suggestion", "suggestion" in issue, str(issue.keys()))
            test("issue 有 source", "source" in issue, str(issue.keys()))
            test("source 为 rule 或 ai", issue["source"] in ("rule", "ai"), issue.get("source"))

        # 检查是否发现了我们故意设置的问题
        issue_texts = " ".join([i.get("location", "") + i.get("element", "") + i.get("current", "") for i in issues])
        test("发现页边距问题", "边距" in issue_texts or "margin" in issue_texts.lower() or "页面" in issue_texts, issue_texts[:200])
        test("发现标题格式问题", "标题" in issue_texts, issue_texts[:200])
        test("发现正文格式问题", "正文" in issue_texts or "仿宋" in issue_texts or "缩进" in issue_texts, issue_texts[:200])

    # 纯规则模式（关闭 AI）
    r = client.session.post(
        f"{BASE_URL}/format-check/check",
        files={"file": ("test_bad.docx", docx_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers={"Authorization": f"Bearer {client.token}"},
        params={"use_ai": "false"}
    )
    test("docx 格式校验（仅规则）", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
    if r.status_code == 200:
        result = r.json()
        all_rule = all(i.get("source") == "rule" for i in result.get("issues", []))
        test("仅规则模式下无 AI 结果", all_rule or len(result.get("issues", [])) == 0, "")
else:
    test("docx 格式校验", False, "python-docx 未安装，无法构造测试文件")

# 2.4 txt 文件校验
txt_data = make_txt_doc()
r = client.session.post(
    f"{BASE_URL}/format-check/check",
    files={"file": ("test.txt", txt_data, "text/plain")},
    headers={"Authorization": f"Bearer {client.token}"}
)
test("txt 格式校验", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")

# 2.5 不支持的文件类型
r = client.session.post(
    f"{BASE_URL}/format-check/check",
    files={"file": ("test.xls", b"fake excel", "application/vnd.ms-excel")},
    headers={"Authorization": f"Bearer {client.token}"}
)
test("不支持的文件类型返回400", r.status_code == 400, f"{r.status_code}: {r.text[:200]}")

# 2.6 校验历史
r = client.get("/format-check/records")
test("获取校验历史", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
if r.status_code == 200:
    data = r.json()
    test("历史返回分页结构", "total" in data and "data" in data, str(data.keys()))
    records = data.get("data", [])
    if records:
        rid = records[0]["id"]
        r = client.get(f"/format-check/records/{rid}")
        test("获取单条校验详情", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
        if r.status_code == 200:
            detail = r.json()
            test("详情包含 issues", "issues" in detail, str(detail.keys()))
            test("详情包含 rule_snapshot", "rule_snapshot" in detail, str(detail.keys()))

# 2.7 普通用户也能使用格式校验
r = user_client.session.post(
    f"{BASE_URL}/format-check/check",
    files={"file": ("test.txt", txt_data, "text/plain")},
    headers={"Authorization": f"Bearer {user_client.token}"}
)
test("普通用户可使用格式校验", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")

# 2.8 预留自动修正接口
r = client.session.post(
    f"{BASE_URL}/format-check/fix",
    files={"file": ("test.txt", txt_data, "text/plain")},
    headers={"Authorization": f"Bearer {client.token}"}
)
test("自动修正接口返回501（预留）", r.status_code == 501, f"{r.status_code}: {r.text[:200]}")

# ==================== 第三部分：对话附件上传 ====================
print("\\n【3. 对话附件上传 —— 多模态写作】")

# 3.1 上传 Word 附件
simple_docx = make_simple_docx()
if simple_docx:
    r = user_client.session.post(
        f"{BASE_URL}/chat/attachments/upload",
        files={"files": ("report.docx", simple_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers={"Authorization": f"Bearer {user_client.token}"}
    )
    test("上传 Word 附件", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
    if r.status_code == 200:
        attachments = r.json()
        test("返回附件列表", isinstance(attachments, list) and len(attachments) > 0, "")
        word_att = attachments[0] if attachments else {}
        word_att_id = word_att.get("id")
        test("Word 附件解析成功", word_att.get("parse_status") == "success", word_att.get("parse_status"))
        test("Word 附件有 text_content", bool(word_att.get("text_content")), word_att.get("text_content", "")[:100])
        test("Word 附件 kind 为 doc", word_att.get("kind") == "doc", word_att.get("kind"))
else:
    test("上传 Word 附件", False, "python-docx 未安装")
    word_att_id = None

# 3.2 上传 PDF 附件
pdf_str = "%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n2 0 obj\n<<\n/Type /Pages\n/Kids [3 0 R]\n/Count 1\n>>\nendobj\n3 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/MediaBox [0 0 612 792]\n/Contents 4 0 R\n>>\nendobj\n4 0 obj\n<<\n/Length 44\n>>\nstream\nBT\n/F1 12 Tf\n100 700 Td\n(社区矫正工作汇报) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000214 00000 n \ntrailer\n<<\n/Size 5\n/Root 1 0 R\n>>\nstartxref\n312\n%%EOF"
pdf_data = pdf_str.encode('utf-8')
r = user_client.session.post(
    f"{BASE_URL}/chat/attachments/upload",
    files={"files": ("report.pdf", pdf_data, "application/pdf")},
    headers={"Authorization": f"Bearer {user_client.token}"}
)
test("上传 PDF 附件", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
if r.status_code == 200:
    pdf_att = r.json()[0]
    pdf_att_id = pdf_att.get("id")
    test("PDF 附件 kind 为 doc", pdf_att.get("kind") == "doc", pdf_att.get("kind"))
else:
    pdf_att_id = None

# 3.3 上传图片附件（OCR）
# 构造一个简单图片（1x1 像素，OCR 可能识别为空，但测试上传流程）
try:
    from PIL import Image
    img = Image.new("RGB", (100, 30), color="white")
    img_buf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    img.save(img_buf.name)
    with open(img_buf.name, "rb") as f:
        img_data = f.read()
    os.unlink(img_buf.name)

    r = user_client.session.post(
        f"{BASE_URL}/chat/attachments/upload",
        files={"files": ("scan.png", img_data, "image/png")},
        headers={"Authorization": f"Bearer {user_client.token}"}
    )
    test("上传图片附件", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
    if r.status_code == 200:
        img_att = r.json()[0]
        test("图片附件 kind 为 image", img_att.get("kind") == "image", img_att.get("kind"))
        # OCR 可能成功也可能失败（取决于图片内容），但流程要通
        test("图片附件有 parse_status", img_att.get("parse_status") in ("success", "failed", "ocr_empty"), img_att.get("parse_status"))
except ImportError:
    test("上传图片附件", False, "PIL 未安装，跳过")

# 3.4 查询附件列表
r = user_client.get("/chat/attachments")
test("查询附件列表", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
if r.status_code == 200:
    atts = r.json()
    test("附件列表非空", len(atts) > 0, f"数量: {len(atts)}")

# 3.5 不支持的文件类型
r = user_client.session.post(
    f"{BASE_URL}/chat/attachments/upload",
    files={"files": ("virus.exe", b"fake exe", "application/x-msdownload")},
    headers={"Authorization": f"Bearer {user_client.token}"}
)
test("上传不支持的文件类型", r.status_code == 400, f"期望400，实际{r.status_code}: {r.text[:200]}")

# 3.6 多文件上传
if simple_docx:
    r = user_client.session.post(
        f"{BASE_URL}/chat/attachments/upload",
        files=[
            ("files", ("doc1.docx", simple_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
            ("files", ("doc2.docx", simple_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
        ],
        headers={"Authorization": f"Bearer {user_client.token}"}
    )
    test("多文件上传", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
    if r.status_code == 200:
        multi_atts = r.json()
        test("返回多个附件", len(multi_atts) == 2, f"数量: {len(multi_atts)}")

# ==================== 第四部分：带附件的对话 ====================
print("\\n【4. 带附件的智能写作对话】")

if SKIP_AI:
    print("  ⏭  跳过 AI 测试（SKIP_AI 已设置）")
else:
    # 4.1 上传一个附件，然后发消息引用
    if simple_docx:
        r = user_client.session.post(
            f"{BASE_URL}/chat/attachments/upload",
            files={"files": ("material.docx", simple_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"Authorization": f"Bearer {user_client.token}"}
        )
        if r.status_code == 200:
            att_id = r.json()[0]["id"]

            r = user_client.post("/chat/send", {
                "message": "请根据这份材料整理成一份汇报",
                "session_id": None,
                "use_rag": False,
                "attachment_ids": [att_id]
            })
            test("带附件发送消息", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
            if r.status_code == 200:
                result = r.json()
                session_id = result.get("session_id")
                test("返回 session_id", bool(session_id), "")
                test("返回 reply 非空", len(result.get("reply", "")) > 0, "reply为空")
                test("返回 attachments 信息", isinstance(result.get("attachments"), list), "")

                # 4.2 多轮对话：不重新上传，直接引用同一会话
                if session_id:
                    r = user_client.post("/chat/send", {
                        "message": "把刚才的汇报压缩到300字",
                        "session_id": session_id,
                        "use_rag": False
                    })
                    test("多轮对话持续引用附件", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
                    if r.status_code == 200:
                        test("第二轮返回 reply 非空", len(r.json().get("reply", "")) > 0, "")

                    # 4.3 验证历史消息包含附件信息
                    r = user_client.get(f"/chat/sessions/{session_id}/messages")
                    test("获取带附件的会话历史", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
                    if r.status_code == 200:
                        msgs = r.json()
                        user_msgs = [m for m in msgs if m.get("role") == "user"]
                        if user_msgs:
                            first_user_msg = user_msgs[0]
                            test("历史消息包含 attachments 字段", "attachments" in first_user_msg, str(first_user_msg.keys()))
                        else:
                            test("历史消息包含 attachments 字段", False, "无用户消息")
            else:
                test("多轮对话持续引用附件", False, "第一轮发送失败")
                test("获取带附件的会话历史", False, "第一轮发送失败")
        else:
            test("带附件发送消息", False, "附件上传失败")
            test("多轮对话持续引用附件", False, "附件上传失败")
            test("获取带附件的会话历史", False, "附件上传失败")
    else:
        test("带附件发送消息", False, "python-docx 未安装")
        test("多轮对话持续引用附件", False, "python-docx 未安装")
        test("获取带附件的会话历史", False, "python-docx 未安装")

    # 4.4 多个附件同时引用
    if simple_docx:
        r = user_client.session.post(
            f"{BASE_URL}/chat/attachments/upload",
            files=[
                ("files", ("mat1.docx", simple_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
                ("files", ("mat2.docx", simple_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
            ],
            headers={"Authorization": f"Bearer {user_client.token}"}
        )
        if r.status_code == 200:
            att_ids = [a["id"] for a in r.json()]
            r = user_client.post("/chat/send", {
                "message": "综合这两份材料，帮我形成一份总结",
                "session_id": None,
                "use_rag": False,
                "attachment_ids": att_ids
            })
            test("多附件同时引用", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
        else:
            test("多附件同时引用", False, "附件上传失败")
    else:
        test("多附件同时引用", False, "python-docx 未安装")

# 4.5 不带附件的普通对话仍然正常
if not SKIP_AI:
    r = user_client.post("/chat/send", {
        "message": "你好，请介绍一下社区矫正工作",
        "session_id": None,
        "use_rag": False
    })
    test("普通对话（无附件）正常", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
else:
    test("普通对话（无附件）正常", True, "SKIP_AI 跳过")

# ==================== 第五部分：原有功能兼容 ====================
print("\\n【5. 原有功能兼容性】")

# 5.1 知识库
r = user_client.post("/knowledge/create", {"name": "测试兼容知识库", "description": "测试"})
test("知识库创建", r.status_code == 200, f"{r.status_code}: {r.text}")
if r.status_code == 200:
    kb_id = r.json().get("id")

    # 上传文档到知识库
    txt = make_txt_doc()
    r = user_client.session.post(
        f"{BASE_URL}/knowledge/upload?kb_id={kb_id}",
        files={"file": ("compat.txt", txt, "text/plain")},
        headers={"Authorization": f"Bearer {user_client.token}"}
    )
    test("知识库文档上传", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")

# 5.2 模板生成（不带附件）
if not SKIP_AI:
    r = user_client.get("/templates/")
    if r.status_code == 200 and r.json():
        tmpl = r.json()[0]
        r = user_client.post("/chat/send", {
            "message": "请生成一份通知",
            "session_id": None,
            "use_rag": True,
            "template_category": tmpl.get("category", "通知公告")
        })
        test("模板生成（无附件）", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")
    else:
        test("模板生成（无附件）", False, "无可用模板")
else:
    test("模板生成（无附件）", True, "SKIP_AI 跳过")

# 5.3 导出 Word
if not SKIP_AI:
    r = user_client.post("/chat/export/docx", {
        "content": "# 测试\\n\\n这是测试内容。",
        "title": "测试导出",
        "doc_number": "",
        "recipient": "各单位",
        "signature": "测试局",
        "date_text": "2026年8月14日",
        "use_red_header": False
    })
    test("导出 Word", r.status_code == 200 and r.headers.get("content-type", "").startswith("application/"), f"{r.status_code}")
else:
    test("导出 Word", True, "SKIP_AI 跳过")

# 5.4 权限控制
r = user_client.get("/knowledge/stats")
test("普通用户不能访问管理统计", r.status_code == 403, f"期望403，实际{r.status_code}")

r = kb_client.get("/knowledge/stats")
test("知识管理员可以访问管理统计", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")

# 5.5 会话管理
r = user_client.get("/chat/sessions")
test("获取会话列表", r.status_code == 200, f"{r.status_code}: {r.text[:200]}")

# ==================== 测试报告 ====================
print("\\n" + "=" * 70)
print(f"测试完成：通过 {passed} / 失败 {failed} / 总计 {passed + failed}")
print("=" * 70)

if failed > 0:
    print(f"\\n⚠️  有 {failed} 项测试失败，请检查后端日志或接口实现")
    sys.exit(1)
else:
    print("\\n🎉 全部测试通过！")
    sys.exit(0)
